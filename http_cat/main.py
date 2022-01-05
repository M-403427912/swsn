print("run")
import json

import requests
import requests
import uuid
import hashlib

import time
import requests
import json
import execjs  # 必须，需要先用pip 安装，用来执行js脚本
from urllib.parse import quote

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt #磅数
from docx.oxml.ns import qn #中文格式
from docx.shared import Inches #图片尺寸
import os

def tran(s):
    # 1.创建百度翻译api需要的数据格式
    input_key = s
    # 实现中英文互译
    data = {
        'q': input_key,
        'from': 'auto',
        'to': 'zh',
        # 百度翻译开发者平台申请
        'appid': '20180613000175689',
        # 利用python标准库生成标准格式的salt数据
        'salt': str(uuid.uuid4()),
    }
    # 按指定要求生成签名数据
    sign = hashlib.md5((data['appid'] +
                        data['q'] + data['salt'] +
                        'rE5J15nPRBF4c8PciThQ').encode('utf-8'))
    # 将bytes类型转化成16进制数据
    sign = sign.hexdigest()
    data['sign'] = sign

    # 2.向百度翻译开发者接口url发送数据,利用post
    url = 'http://api.fanyi.baidu.com/api/trans/vip/translate'
    # 利用requests的特点:即使url或者data中格式即使不是标准
    # urlencoded类型, 也可以完成自动转换.
    # 永远不要忘记一个合格spider的自我修养: 加user-agent
    headers = {
        'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/80.0.3987.149 Safari/537.36',
    }
    response = requests.post(url, headers=headers, data=data)
    response_json = response.json()

    # 3.处理接收到的json数据
    result_list = response_json['trans_result']
    result_str = ''
    result_len = len(result_list) - 1
    for index, each in enumerate(result_list):
        result_str += each['dst']
        if index != result_len:
            result_str += '\n'
    return('{}'.format(result_str))



headers = {
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/96.0.4664.110 Safari/537.36',
    'Referer': 'https://cat.benesse.ne.jp/gallery/',
    'Connection': 'keep-alive',
    'content-type': 'application/json',
    'sec-ch-ua-mobile':"0",
    'sec-ch-ua-platform': "Windows",
    'sec-fetch-dest': 'empty',
    'sec-fetch-mode': 'cors',
    'sec-fetch-site': 'same-origin'
}

timeOut = 30


def download_img(img_name, img_url):
    # header = {"Authorization": "Bearer " + api_token}  # 设置http header
    #print(img_url)

    try:
        response = requests.get(img_url,timeout=timeOut)
        filename = "png/" + img_name + ".png"
        #print(response)
        if (response.status_code  == 200):
            with open(filename, "wb") as f:
                f.write(response.content)  # 将内容写入图片

    except BaseException:
        print(filename +" 下载失败")


HTTPHEADSTR = "https://cat.benesse.ne.jp"
HTTPJSONDATA = "https://cat.benesse.ne.jp/gallery/api/v1/photo?query=query+FetchPhoto{photo{id,pet_type,sex,pet_name,pet_kind,pic_age,comment,image}}"

payload = {  # 定义参数
    'operationName': "FetchPhoto",
    'query': 'query FetchPhoto($id: Int, $page: Int, $pet_kind: Int, $theme_code: Int, $pet_type: Int, $sex: Int, $keyword: String) {\n  photo(id: $id, page: $page, pet_kind: $pet_kind, theme_code: $theme_code, pet_type: $pet_type, sex: $sex, keyword: $keyword) {\n    id\n    post_id\n    pet_type\n    sex\n    pet_name\n    pet_kind\n    pet_kind_code\n    pic_age\n    theme_code\n    theme_name\n    comment\n    image\n    photo_type\n    __typename\n  }\n}\n',
    'variables': {
        'id': None,
        'page': 1,
        'pet_type': 2,
        'sex': None
    }
}

response = requests.post(HTTPJSONDATA, data=json.dumps(payload), headers=headers,timeout=timeOut)


data = response.content.decode('utf-8')
# print(data)
jsonData = json.loads(data)
# print(jsonData)



dataDict = dict()

for photo in jsonData["data"]["photo"]:
    photo["isnew"]=True;
    if os.path.exists('png/'+str(photo["id"])+'.png'):
        print(str(photo["id"])+" 已存在")
        photo["isnew"]=False;
        continue;

    sex = "♂" if photo["sex"]==1 else "♀"
    comment="";
    if photo["comment"]!=None:
        comment=tran(str(photo["comment"]or '')).replace("\n","")
    else:
        photo["isnew"]=False;
        continue



    photo["comment"]=comment;
    print("id："+str(photo["id"])+"\t介绍："+comment)
    download_img(str(photo["id"]), HTTPHEADSTR + photo["image"][0])
    time.sleep(1.1)


today = time.strftime("%Y{y}%m{m}%d{d}",time.localtime()).format(y='年',m='月',d='日')

for photo in jsonData["data"]["photo"]:
    if photo["isnew"]:
        document = Document()
        # 设置文档的基础字体中文
        document.styles['Normal'].font.name = u'宋体'
        # 设置文档的基础样式
        document.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        pl = document.add_paragraph()
        # 对齐方式为居中，没有这句话默认左对齐
        pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = pl.add_run("")

        run.add_picture('png/'+str(photo["id"])+'.png',)
        # 初始化建立第一个自然段
        pl = document.add_paragraph()
        # 对齐方式为居中，没有这句话默认左对齐
        pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pl.add_run(photo["comment"])
        document.add_paragraph()
        document.save(photo["comment"]+'.docx')


print("over")

